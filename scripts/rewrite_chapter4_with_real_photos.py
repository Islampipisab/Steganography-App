from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


SOURCE_DOC = r"C:\Users\pipis\Downloads\CSE GradII_stego_Chapters-ch4-5 (1).docx"
OUTPUT_DOC = r"C:\Users\pipis\Downloads\CSE GradII_stego_Chapters-ch4-5_CH4_REAL_PHOTOS.docx"

IMG_HOME = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-home.png"
IMG_ENCODE_OK = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-encode-success.png"
IMG_DECODE = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-decode-tab.png"
IMG_HIDE_TEXT = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-hide-text-tab.png"
IMG_HIDE_BARCODE = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-hide-barcode-tab.png"
IMG_EXTRACT = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-extract-tab.png"


def qn(tag):
    prefix, local = tag.split(":")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return "{%s}%s" % (ns[prefix], local)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def insert_text(anchor, text, style="Normal", bold=False):
    p = anchor.insert_paragraph_before(text, style=style)
    if bold and p.runs:
        p.runs[0].bold = True
    return p


def insert_code(anchor, code):
    p = anchor.insert_paragraph_before("", style="Normal")
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F3F3")
    p._element.get_or_add_pPr().append(shd)


def insert_image(anchor, image_path, width=5.8):
    p = anchor.insert_paragraph_before("", style="Normal")
    p.add_run().add_picture(image_path, width=Inches(width))


def main():
    doc = Document(SOURCE_DOC)

    ch5_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().lower() == "chapter v":
            ch5_idx = i
            break
    if ch5_idx is None:
        raise RuntimeError("Could not find Chapter V in thesis.")

    anchor = doc.paragraphs[ch5_idx]

    entries = [
        ("text", "4.1 Development Environment", "Title"),
        ("text", "The implementation was developed using Python 3 with Flask for the web backend, NumPy and Pillow for image processing, and cryptographic modules for secure payload protection. Deployment was validated on a Linux cloud server using Gunicorn with systemd service management.", "Normal"),
        ("text", "Table 4.1 Software tools and libraries used in system implementation: Python, Flask, NumPy, Pillow, Gunicorn, systemd, nginx, AES/RSA cryptography modules.", "Normal"),

        ("text", "4.2 System Architecture Implementation", "Title"),
        ("text", "The implemented architecture is modular: config, core steganography methods (LSB/DCT/Hybrid), barcode encode/decode module, crypto modules, metrics module, processing/denoising, and web UI layer. This modular design allows each function to be tested independently and integrated into an end-to-end secure pipeline.", "Normal"),

        ("text", "4.3 Encryption Modules", "Title"),
        ("text", "4.3.1 SecureEncryption Module", "Title"),
        ("text", "The SecureEncryption path is responsible for payload protection before embedding. It transforms plaintext to ciphertext to reduce information leakage risk in case hidden data is partially exposed.", "Normal"),
        ("text", "4.3.2 HybridCrypto Module", "Title"),
        ("text", "A hybrid model combines symmetric encryption for payload efficiency and asymmetric encryption for secure key handling. This increases confidentiality and practical deployability.", "Normal"),

        ("text", "4.4 Custom Barcode Carrier Generation", "Title"),
        ("text", "The payload is encoded into a custom barcode-like binary pattern. The encoder embeds header, length information, and protected data bits into a structured grid before steganographic hiding.", "Normal"),
        ("code", "@app.route(\"/api/encode\", methods=[\"POST\"])\ndef api_encode():\n    cbs = get_cbs()\n    pattern = cbs.encode(text, use_hybrid_encryption=use_hybrid)\n    return send_file(_barcode_pattern_to_bytes(pattern), mimetype=\"image/png\")", "Normal"),
        ("text_bold", "Figure 4.1 Main application interface of the steganography system.", "Normal"),
        ("image", IMG_HOME, "Normal"),
        ("text_bold", "Figure 4.2 Encode tab showing successful barcode generation and download.", "Normal"),
        ("image", IMG_ENCODE_OK, "Normal"),

        ("text", "4.5 Barcode Decoding", "Title"),
        ("text", "The decode pipeline loads barcode images, extracts encoded bits, validates format headers, and reconstructs the original payload. Error handling reports invalid format or corrupted patterns.", "Normal"),
        ("text_bold", "Figure 4.3 Decode tab for barcode recovery workflow.", "Normal"),
        ("image", IMG_DECODE, "Normal"),

        ("text", "4.6 Hiding Existing Barcode (Embedding Pre-Generated Barcode)", "Title"),
        ("text", "4.6.1 Workflow", "Title"),
        ("text", "A pre-generated barcode is loaded and embedded into a selected cover image using LSB, DCT, or Hybrid mode. The selected method defines capacity, robustness, and imperceptibility trade-offs.", "Normal"),
        ("text_bold", "Figure 4.4 Hide Text tab in the implemented web interface.", "Normal"),
        ("image", IMG_HIDE_TEXT, "Normal"),
        ("text_bold", "Figure 4.5 Hide Barcode tab for embedding pre-generated barcode carriers.", "Normal"),
        ("image", IMG_HIDE_BARCODE, "Normal"),

        ("text", "4.7 Steganography Embedding, Extraction, and Quality Evaluation", "Title"),
        ("text", "4.7.1 LSB-Based Steganography", "Title"),
        ("text", "LSB embedding provides high payload capacity with minimal visual distortion in lossless scenarios.", "Normal"),
        ("text", "4.7.2 Data Extraction Process", "Title"),
        ("text", "Extraction recovers hidden bits from stego images, reconstructs barcode/text payloads, and returns decoded content through API endpoints.", "Normal"),
        ("text", "4.7.3 DCT-Based Steganography", "Title"),
        ("text", "DCT embedding modifies selected frequency coefficients for improved robustness under moderate transformations.", "Normal"),
        ("text", "4.7.4 Hybrid Embedding Strategy", "Title"),
        ("text", "Hybrid mode combines LSB and DCT domains to balance visual quality, robustness, and extraction reliability.", "Normal"),
        ("text", "4.7.5 Quality Metrics Evaluation", "Title"),
        ("text", "Quality is assessed using PSNR, SSIM, BER, entropy change, and related comparisons between cover and stego images.", "Normal"),

        ("text", "4.8 Noise Handling and Denoising", "Title"),
        ("text", "Median and Gaussian filters are integrated to improve extraction reliability after noise corruption. This supports practical robustness testing in non-ideal transmission conditions.", "Normal"),
        ("text", "4.8.1 Effect of DCT Embedding and Noise on Quality Metrics", "Title"),
        ("text", "The implementation tracks metric degradation under noisy conditions and evaluates denoising effectiveness before extraction.", "Normal"),

        ("text", "4.9 Graphical User Interface", "Title"),
        ("text", "The GUI provides interactive tabs for encode, decode, hide text, hide barcode, and extract operations with progress bars and detailed logs.", "Normal"),
        ("text_bold", "Figure 4.6 Extract tab for recovering data from stego images.", "Normal"),
        ("image", IMG_EXTRACT, "Normal"),

        ("text", "4.10 End-to-End Execution Flow", "Title"),
        ("text", "4.10.1 Sender Side", "Title"),
        ("text", "Input -> optional encryption -> barcode generation -> cover selection -> embedding -> stego output.", "Normal"),
        ("text", "4.10.2 Receiver Side", "Title"),
        ("text", "Stego input -> extraction -> barcode decode -> optional decryption -> recovered message/image.", "Normal"),

        ("text", "4.11 Summary", "Title"),
        ("text", "This chapter presented the full preliminary implementation of the secure image steganography system, validated using live successful application workflows and interface captures. The next chapter focuses on testing and performance evaluation.", "Normal"),
    ]

    for kind, value, style in entries:
        if kind == "text":
            insert_text(anchor, value, style=style)
        elif kind == "text_bold":
            insert_text(anchor, value, style=style, bold=True)
        elif kind == "code":
            insert_code(anchor, value)
        elif kind == "image":
            insert_image(anchor, value)

    # Remove existing Chapter IV body while preserving the top title lines and Chapter V onward.
    old_section = list(doc.paragraphs[2:ch5_idx])
    for p in old_section:
        delete_paragraph(p)

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
