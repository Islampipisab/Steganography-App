"""
One-off: fill Chapter V sections 5.1–5.5 in the thesis docx (keeps all existing headings).
Run: python scripts/fill_chapter5_docx.py
"""
from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def main() -> None:
    path = r"c:\Users\pipis\Downloads\CSE GradII_stego_Chapters-ch4-5 (2).docx"
    doc = Document(path)

    # --- 5.1 Introduction (expand first paragraph after heading)
    intro_heading = "5.1 Introduction"
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == intro_heading:
            if i + 1 < len(doc.paragraphs):
                p2 = doc.paragraphs[i + 1]
                p2.text = (
                    "This chapter presents the testing and evaluation of the proposed steganography system. "
                    "The objective is to verify that the system meets the functional and non-functional "
                    "requirements defined in Chapter 3. The testing process includes automated verification "
                    "of the Flask web API (local test client), live deployment checks against the public HTTPS "
                    "endpoint (stegax.design), manual and browser-based smoke tests, and performance evaluation "
                    "using PSNR, SSIM, and BER where applicable. Field testing on mobile (iPhone) was also used "
                    "to validate real-user workflows (save/share stego images, long-running extract operations)."
                )
                insert_paragraph_after(
                    p2,
                    "Testing covers: (1) barcode encode/decode with optional hybrid (AES+RSA) encryption; "
                    "(2) hide text and hide barcode using LSB, DCT, and Hybrid embedding; (3) symmetric "
                    "extraction (same method as embedding); (4) quality metrics API; (5) service reliability "
                    "under concurrent HTTP traffic when long-running jobs execute (Gunicorn worker model).",
                )
            break

    # Replace known placeholder lines by exact match
    replacements: dict[str, str] = {
        "Hardware specifications": (
            "Development: PC with Windows 10/11, 8–16 GB RAM, SSD; remote SSH to DigitalOcean droplet "
            "(Ubuntu, Python 3.12, venv). Deployment server: cloud VPS at 46.101.244.129, hostname Stegax, "
            "domain stegax.design behind Nginx reverse proxy (TLS) to Gunicorn on 127.0.0.1:5000."
        ),
        "Software environment": (
            "Python 3.x; Flask web application; Gunicorn with gthread worker (1 worker, 8 threads) and "
            "600 s timeout; Nginx with extended proxy read/send timeouts and large client body limit for "
            "image uploads; systemd unit stego-web.service for automatic restarts."
        ),
        "Libraries used": (
            "Pillow, NumPy, SciPy (DCT), scikit-image, cryptography (Flask app requirements as pinned in "
            "requirements.txt). Desktop Tkinter UI optional; primary validation path is the web UI + REST API."
        ),
        "Different image types (JPEG, PNG)": (
            "PNG and BMP for lossless barcode carriers and stego I/O; JPEG accepted where the pipeline allows "
            "(DCT path supports JPEG with robust mode). TIFF supported for extract/upload in the web app."
        ),
        "Different resolutions": (
            "Covers from small test grids (e.g. 256×256, 320×320, 384×384) up to larger photos; barcode "
            "pattern scaled to cover dimensions in hide-barcode. Large patterns trigger omission of heavy "
            "JSON previews to protect mobile browsers."
        ),
        "Various message sizes": (
            "Short ASCII strings through longer payloads; hybrid-encrypted barcodes; stress test of repeated "
            "extract-barcode (20 consecutive API calls on the same stego PNG) to detect intermittent failures."
        ),
        "Each module is tested independently:": (
            "Each module is tested independently (automated script scripts/run_app_verification.py — 9/9 checks):"
        ),
        "AES encryption module": (
            "AES encryption module — AES stage of the hybrid pipeline exercised via Encode with “Use hybrid "
            "cryptography” and Decode with hybrid decryption enabled in the barcode engine."
        ),
        "RSA module": (
            "RSA module — Covered as part of the hybrid encode/decode round-trip (RSA-wrapped AES key material)."
        ),
        "QR generation module": (
            "QR generation module — Custom barcode (QR-like grid) generation and rendering to PNG via /api/encode; "
            "load/decode via /api/decode and pattern-based extract paths."
        ),
        "LSB embedding": (
            "LSB embedding — LSB hide-text and hide-barcode; vectorized bit packing/extraction for performance; "
            "PNG output with controlled compression level for barcode stego."
        ),
        "DCT embedding": (
            "DCT embedding — DCT hide-text and hide-barcode; optional JPEG-robust coefficient quantization; strict "
            "DCTBAR header matching on lossless PNG after field bugfix (prevents false “success” when wrong extract "
            "method is chosen)."
        ),
        "Encryption + QR encoding": (
            "Form POST /api/encode → PNG barcode download; optional X-Progress-Log header; JSON legacy path "
            "retained with ?format=json for tests."
        ),
        "QR + embedding": (
            "Multipart /api/hide-text and /api/hide-barcode with cover + payload; hide-barcode default response "
            "raw PNG with X-Metrics-Json / X-Progress-Log headers (avoids huge base64-in-JSON failures)."
        ),
        "Embedding + extraction": (
            "Round-trip for each method (LSB, DCT, Hybrid): hide then extract-text or extract-barcode with "
            "matching method; live HTTPS tests to stegax.design confirmed LSB/DCT/Hybrid barcode pipeline."
        ),
        "End-to-end system testing": (
            "End-to-end system testing: browser smoke test on stegax.design (Encode → Generate barcode succeeded); "
            "API-level hide-barcode + extract-barcode for all three methods with shared test assets; desktop "
            "verification suite passes in full."
        ),
        "Verify correct message recovery": (
            "Verify correct message recovery: extracted text matches original payload for text and barcode-decode "
            "modes when embedding and extraction methods match; mismatch (e.g. LSB stego extracted as DCT) yields "
            "errors or garbage — documented as user-configuration risk mitigated by UI hints."
        ),
        "Verify image quality preservation": (
            "Verify image quality preservation: PSNR/SSIM/BER from /api/metrics and hide-barcode response headers "
            "when output size/pixel count within server thresholds; skipped for very large outputs to avoid stalls."
        ),
        "Full pipeline execution": (
            "Full pipeline execution on production stack: client → Nginx (443) → Gunicorn (127.0.0.1:5000) → "
            "Flask app; systemd-managed service restart after code updates; pip install -r requirements.txt in venv."
        ),
        "Correct extraction": (
            "Correct extraction validated for aligned method selection; Decode tab must receive the barcode PNG "
            "from Encode — uploading a stego image to Decode produces “Invalid barcode format” (observed on mobile)."
        ),
        "Data integrity": (
            "Data integrity: lossless PNG roundtrip preserves embedded barcode bits for LSB/DCT/Hybrid when the "
            "image is not recompressed (iPhone: re-saving stego as JPEG can harm DCT unless JPEG-robust hiding was used)."
        ),
        "System reliability": (
            "System reliability: replaced single synchronous Gunicorn worker with gthread + threads so long "
            "hide/extract jobs no longer block the entire site; extended client, proxy, and JavaScript timeouts "
            "to reduce “failed to fetch” on mobile."
        ),
        "Compare different techniques:": (
            "Compare different techniques (LSB vs DCT vs Hybrid) on imperceptibility, robustness, and measured "
            "metrics where available; field issues below are included in the test record."
        ),
    }

    for p in doc.paragraphs:
        key = p.text.strip()
        if key in replacements:
            p.text = replacements[key]

    # Block after "Compare different techniques" — find paragraph and insert error log
    for p in doc.paragraphs:
        if p.text.startswith("Compare different techniques") and "field issues" in p.text:
            q = insert_paragraph_after(
                p,
                "Recorded defects and resolutions during testing: (1) Slow or failing hide-barcode / extract — "
                "mitigated by vectorized NumPy operations in LSB/DCT/Hybrid barcode paths and binary PNG responses. "
                "(2) Browser/proxy timeouts and “failed to fetch” on long jobs — mitigated by Gunicorn --timeout 600, "
                "Nginx proxy timeouts, 10-minute fetch in web JS, gthread worker. (3) Entire site unresponsive during "
                "one long request — mitigated by gthread + multi-threaded worker instead of one blocking sync worker. "
                "(4) iPhone: DCT extract showed garbage while LSB worked — root cause: extraction method must match "
                "embedding method; fuzzy DCT header on lossless images could false-match wrong stego; fixed by exact "
                "header for PNG/BMP/TIFF; UI hints added under Hide Barcode and Extract. (5) Decode tab error on "
                "mobile — users uploaded stego instead of raw barcode image. (6) IDE browser automation could not "
                "set files on hidden <input type=\"file\"> (not visible to Playwright); validation used API instead. "
                "(7) Deployment: no SSH key in cloud IDE — upload must run deploy.ps1 from Windows; on server, "
                "bash-only — attempts to run .ps1 or cd C:\\... failed until Linux commands (systemctl restart "
                "stego-web) were used. (8) On Linux servers use curl, not curl.exe.",
            )
            insert_paragraph_after(
                q,
                "Automated regression: scripts/run_app_verification.py reports 9/9 passed (encode/decode, "
                "hide/extract text LSB+DCT+Hybrid, hide/extract barcode LSB+DCT+Hybrid, 20× LSB extract stress, "
                "metrics API).",
            )
            break

    # Fill comparison table (table 0)
    tbl = doc.tables[0]
    # Row 0 is header
    data = [
        (
            "LSB",
            "Typically highest among the three on lossless PNG (minimal pixel change).",
            "High structural similarity to cover when capacity allows.",
            "Very low BER on lossless roundtrip; rises if image is recompressed or edited.",
            "Weakest to lossy JPEG/social recompression; best for pristine PNG workflow.",
        ),
        (
            "DCT",
            "Moderate; coefficient quantization trades imperceptibility for robustness options.",
            "Good when blocks align; JPEG-robust mode targets recompression survival.",
            "Low on lossless roundtrip; normal DCT hiding is sensitive to JPEG if not using JPEG-robust.",
            "Better than pure LSB against JPEG when JPEG-robust embedding is enabled.",
        ),
        (
            "Hybrid",
            "Intermediate—splits payload between spatial LSB and frequency DCT channels.",
            "Generally good; reflects combined spatial/frequency perturbation.",
            "Low on lossless roundtrip with matched extraction; method must match embed.",
            "Combines aspects of LSB and DCT; user must keep extract method aligned with embed.",
        ),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            tbl.rows[r].cells[c].text = val

    doc.save(path)
    print("Saved:", path)


if __name__ == "__main__":
    main()
