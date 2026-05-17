from docx import Document

SRC = r"C:\Users\pipis\Downloads\CSE_GradII_Chapter4_Only_Detailed_Final.docx"
OUT = r"C:\Users\pipis\Downloads\CSE_GradII_Chapter4_Implementation_Final.docx"


intro_heading = "4.1 Introduction"
intro_p1 = (
    "This chapter presents the implementation phase of the proposed secure steganography system. "
    "Based on the architecture, data flow, and algorithmic design established in Chapter 3, the implemented system "
    "translates the conceptual model into a practical and testable software solution."
)
intro_p2 = (
    "The implementation covers barcode generation, cryptography integration, LSB/DCT/Hybrid embedding workflows, "
    "extraction and decoding procedures, runtime deployment considerations, and user-interface execution paths. "
    "Accordingly, this chapter acts as the bridge between design decisions in Chapter 3 and the measurable testing "
    "results discussed in the following chapter."
)


def main():
    d = Document(SRC)

    # Rename chapter subtitle if needed.
    for p in d.paragraphs:
        t = (p.text or "").strip().lower()
        if t == "preliminary implementation":
            p.text = "Implementation"
            break

    # Replace existing first section heading to Introduction.
    first_sec_idx = None
    for i, p in enumerate(d.paragraphs):
        if (p.text or "").strip().startswith("4.1 "):
            first_sec_idx = i
            break

    if first_sec_idx is not None:
        d.paragraphs[first_sec_idx].text = intro_heading

        # Remove immediate old intro/body lines before next numbered section.
        j = first_sec_idx + 1
        to_remove = []
        while j < len(d.paragraphs):
            txt = (d.paragraphs[j].text or "").strip()
            if txt.startswith("4.2 "):
                break
            if txt:
                to_remove.append(d.paragraphs[j])
            j += 1
        for p in to_remove:
            el = p._element
            el.getparent().remove(el)

        # Insert new detailed intro right before 4.2.
        next_anchor = None
        for p in d.paragraphs:
            if (p.text or "").strip().startswith("4.2 "):
                next_anchor = p
                break
        if next_anchor is not None:
            next_anchor.insert_paragraph_before(intro_p2)
            next_anchor.insert_paragraph_before(intro_p1)

    d.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

