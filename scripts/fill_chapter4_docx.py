from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

SOURCE_DOC = r"C:\Users\pipis\Downloads\CSE GradII_stego_Chapters-ch4-5 (1).docx"
OUTPUT_DOC = r"C:\Users\pipis\Downloads\CSE GradII_stego_Chapters-ch4-5_filled.docx"
SCREENSHOT = r"C:\Users\pipis\.cursor\projects\c-Users-pipis-OneDrive-Desktop-newgrad\assets\c__Users_pipis_AppData_Roaming_Cursor_User_workspaceStorage_73701bacba3b1b782b5361dccfd68ec7_images_image-34b21076-a39a-4446-bd60-514370200c7e.png"


def qn(tag):
    prefix, local = tag.split(":")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return "{%s}%s" % (ns[prefix], local)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def insert_text(anchor, text, style="Normal", bold=False):
    p = anchor.insert_paragraph_before(text, style=style)
    if bold and p.runs:
        p.runs[0].bold = True


def insert_code(anchor, code_text):
    p = anchor.insert_paragraph_before("", style="Normal")
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F3F3")
    p._element.get_or_add_pPr().append(shd)


def insert_image(anchor, image_path, width_in=5.8):
    p = anchor.insert_paragraph_before("", style="Normal")
    p.add_run().add_picture(image_path, width=Inches(width_in))


def main():
    doc = Document(SOURCE_DOC)

    ch5_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().lower() == "chapter v":
            ch5_idx = i
            break
    if ch5_idx is None:
        raise RuntimeError("Chapter V not found.")

    anchor = doc.paragraphs[ch5_idx]

    entries = [
        ("text", "4.1 Introduction", "Title"),
        ("text", "This chapter presents the practical implementation of the secure steganography system. The implementation converts the conceptual framework into working software that supports custom barcode generation, optional hybrid cryptography, multiple embedding strategies (LSB, DCT, Hybrid), and a browser-based interface for end users.", "Normal"),

        ("text", "4.2 Hardware and Deployment Implementation", "Title"),
        ("text", "4.2.1 Development Hardware", "Title"),
        ("text", "Development was performed on a standard personal computer with multi-core CPU, 8+ GB RAM, and SSD storage.", "Normal"),
        ("text", "4.2.2 Cloud Deployment Hardware", "Title"),
        ("text", "The system was deployed on a cloud Linux server (DigitalOcean droplet), configured for 24/7 operation with systemd and remote access.", "Normal"),

        ("text", "4.3 Software Implementation", "Title"),
        ("text", "4.3.1 Development Tools and Libraries", "Title"),
        ("text", "Programming language: Python 3", "Normal"),
        ("text", "Web framework: Flask", "Normal"),
        ("text", "Image processing: Pillow, NumPy", "Normal"),
        ("text", "Cryptography: AES and RSA hybrid model", "Normal"),
        ("text", "Deployment runtime: Gunicorn + systemd (+ nginx for reverse proxy/SSL)", "Normal"),

        ("text", "4.3.2 Web API Endpoints", "Title"),
        ("text", "The backend provides dedicated APIs for encoding, hiding, extraction, decoding, denoising, and quality metrics.", "Normal"),
        ("code", "@app.route(\"/api/encode\", methods=[\"POST\"])\ndef api_encode():\n    pattern = get_cbs().encode(text, use_hybrid_encryption=use_hybrid)\n    return send_file(_barcode_pattern_to_bytes(pattern), mimetype=\"image/png\")"),

        ("text", "4.3.3 Core Embedding Methods", "Title"),
        ("text", "(a) LSB method: high capacity and high visual quality for lossless scenarios.", "Normal"),
        ("text", "(b) DCT method: improved robustness for compressed-image workflows.", "Normal"),
        ("text", "(c) Hybrid LSB-DCT method: balances capacity, imperceptibility, and robustness.", "Normal"),
        ("code", "if method == \"LSB\":\n    get_barcode_stego().hide_barcode_in_image(cover_path, pattern, out_path)\nelif method == \"DCT\":\n    get_dct().hide_barcode_in_image(cover_path, pattern, out_path)\nelse:\n    get_hybrid().hide_barcode_in_image(cover_path, pattern, out_path)"),

        ("text", "4.3.4 Extraction and Decoding Logic", "Title"),
        ("text", "Extraction supports text-only or barcode-decode paths. A reliability enhancement is used for scaled barcode patterns by retrying decode using downsampled variants.", "Normal"),
        ("code", "def _decode_barcode_pattern(cbs, pattern, use_hybrid_decryption=True):\n    try:\n        return cbs.decode(pattern, use_hybrid_decryption=use_hybrid_decryption)\n    except ValueError:\n        for k in range(2, min(pattern.shape) // 3 + 1):\n            small = pattern[::k, ::k].copy()\n            try:\n                return cbs.decode(small, use_hybrid_decryption=use_hybrid_decryption)\n            except ValueError:\n                pass\n    raise"),

        ("text", "4.3.5 Quality and Robustness Evaluation Integration", "Title"),
        ("text", "The implementation includes quality analysis panels and APIs reporting PSNR, SSIM, BER, entropy, and pixel-level differences.", "Normal"),

        ("text", "4.3.6 User Interface Implementation", "Title"),
        ("text", "The web interface is organized into tabs: Encode, Decode, Hide Text, Hide Barcode, and Extract. It includes operation logs, progress bars, file controls, and result previews.", "Normal"),
        ("text_bold", "Figure 4.1 Program screenshot (Extract from stego image workflow)", "Normal"),
        ("image", SCREENSHOT, "Normal"),

        ("text", "4.4 System Integration", "Title"),
        ("text", "4.4.1 Module Integration", "Title"),
        ("text", "Pipeline: input data -> optional encryption -> barcode generation -> stego embedding -> extraction -> decoding -> optional decryption.", "Normal"),
        ("text", "4.4.2 End-to-End Execution Flow", "Title"),
        ("text", "Sender side: input text/image -> optional AES+RSA encryption -> barcode generation -> hide in cover image -> save stego image.", "Normal"),
        ("text", "Receiver side: upload stego image -> extract bits/pattern -> decode barcode -> optional hybrid decryption -> recover original content.", "Normal"),

        ("text", "4.5 Summary", "Title"),
        ("text", "This chapter presented the implementation of the proposed steganography system, including backend services, core embedding methods, cryptographic protection, and the web user interface. The system is modular and deployment-ready; Chapter 5 presents testing and performance evaluation.", "Normal"),
    ]

    for entry in entries:
        kind = entry[0]
        value = entry[1]
        style = entry[2] if len(entry) > 2 else "Normal"
        if kind == "text":
            insert_text(anchor, value, style=style)
        elif kind == "text_bold":
            insert_text(anchor, value, style=style, bold=True)
        elif kind == "code":
            insert_code(anchor, value)
        elif kind == "image":
            insert_image(anchor, value)

    old = list(doc.paragraphs[2:ch5_idx])
    for p in old:
        delete_paragraph(p)

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
