from pathlib import Path
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


DOC_PATH = Path(r"C:\Users\pipis\Downloads\CSE GradII_stego_Chapters-ch4-5 (1).docx")
BACKUP_PATH = DOC_PATH.with_name("CSE GradII_stego_Chapters-ch4-5 (1)_backup_before_ch4_update.docx")

IMG_HOME = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-home.png"
IMG_ENCODE_OK = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-encode-success.png"
IMG_DECODE = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-decode-tab.png"
IMG_HIDE_TEXT = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-hide-text-tab.png"
IMG_HIDE_BARCODE = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-hide-barcode-tab.png"
IMG_EXTRACT = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-extract-tab.png"


def qn(tag):
    prefix, local = tag.split(":")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return "{%s}%s" % (ns[prefix], local)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def insert_text(anchor, text, style="Normal", bold=False):
    p = anchor.insert_paragraph_before(text, style=style)
    if bold and p.runs:
        p.runs[0].bold = True


def insert_code(anchor, code_text):
    p = anchor.insert_paragraph_before("", style="Normal")
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F3F3F3")
    p._element.get_or_add_pPr().append(shade)


def insert_image(anchor, image_path):
    p = anchor.insert_paragraph_before("", style="Normal")
    p.add_run().add_picture(image_path, width=Inches(5.8))


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(str(DOC_PATH))

    ch5_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().lower() == "chapter v":
            ch5_idx = i
            break
    if ch5_idx is None:
        raise RuntimeError("Could not find Chapter V.")

    anchor = doc.paragraphs[ch5_idx]

    entries = [
        ("text", "4.1 Development Environment", "Title"),
        ("text", "The system was implemented using Python 3, Flask, NumPy, and Pillow. The application was deployed on a Linux cloud server using Gunicorn and systemd for 24/7 operation.", "Normal"),
        ("code", "app = Flask(__name__, template_folder=\"templates\", static_folder=\"static\")\napp.config[\"MAX_CONTENT_LENGTH\"] = 200 * 1024 * 1024  # 200 MB", "Normal"),
        ("text", "Explanation: The web application is initialized and configured to support large image uploads required by steganography operations.", "Normal"),

        ("text", "4.2 System Architecture Implementation", "Title"),
        ("text", "The architecture is modular and includes core steganography methods, barcode encoding/decoding, cryptography, metrics, denoising, and web UI modules.", "Normal"),
        ("code", "from core import BarcodeSteganography, DCTSteganography, HybridSteganography\nfrom barcode import CustomBarcodeSteganography\nfrom metrics import QualityMetrics", "Normal"),
        ("text", "Explanation: Each module has a clear responsibility, which improves maintainability and allows independent testing.", "Normal"),

        ("text", "4.3 Encryption Modules", "Title"),
        ("text", "4.3.1 SecureEncryption Module", "Title"),
        ("text", "The secure encryption path protects plaintext before embedding into the carrier.", "Normal"),
        ("code", "if use_hybrid_encryption and self.hybrid_crypto:\n    encrypted_result = self.hybrid_crypto.encrypt_text_hybrid(text)\n    data_to_encode = \"[HYBRID_ENCRYPTED]\" + encrypted_result", "Normal"),
        ("text", "Explanation: The plaintext is transformed into protected ciphertext before barcode generation.", "Normal"),
        ("text", "4.3.2 HybridCrypto Module", "Title"),
        ("text", "Hybrid cryptography combines efficient symmetric encryption with secure key exchange.", "Normal"),
        ("code", "private_key, public_key = HybridCrypto.load_rsa_keys()\nself.hybrid_crypto = HybridCrypto(rsa_public_key=public_key, rsa_private_key=private_key)", "Normal"),
        ("text", "Explanation: RSA keys are loaded and used by the hybrid cryptography module to secure encryption/decryption workflows.", "Normal"),

        ("text", "4.4 Custom Barcode Carrier Generation", "Title"),
        ("text", "The payload is converted into a structured barcode-like pattern that can later be hidden in cover images.", "Normal"),
        ("text_bold", "Figure 4.1 Main application interface of the steganography system.", "Normal"),
        ("image", IMG_HOME, "Normal"),
        ("code", "@app.route(\"/api/encode\", methods=[\"POST\"])\ndef api_encode():\n    pattern = get_cbs().encode(text, use_hybrid_encryption=use_hybrid)\n    return send_file(_barcode_pattern_to_bytes(pattern), mimetype=\"image/png\")", "Normal"),
        ("text", "Explanation: This endpoint receives user text, generates the barcode pattern, then returns a PNG barcode image.", "Normal"),

        ("text_bold", "Figure 4.2 Encode tab showing successful barcode generation and download.", "Normal"),
        ("image", IMG_ENCODE_OK, "Normal"),
        ("code", "log.append(\"40% – Encoding to barcode pattern\")\npattern = cbs.encode(text, use_hybrid_encryption=use_hybrid)\nlog.append(\"80% – Rendering PNG\")", "Normal"),
        ("text", "Explanation: The success status appears after pattern generation and rendering complete without errors.", "Normal"),

        ("text", "4.5 Barcode Decoding", "Title"),
        ("text_bold", "Figure 4.3 Decode tab for barcode recovery workflow.", "Normal"),
        ("image", IMG_DECODE, "Normal"),
        ("code", "pattern = cbs.load_barcode(path)\ntext = cbs.decode(pattern, use_hybrid_decryption=True)\nreturn jsonify({\"text\": text, \"is_image\": False})", "Normal"),
        ("text", "Explanation: The decoder loads the barcode image, reconstructs payload bits, and returns decoded content.", "Normal"),

        ("text", "4.6 Hiding Existing Barcode (Embedding Pre-Generated Barcode)", "Title"),
        ("text", "4.6.1 Workflow", "Title"),
        ("text_bold", "Figure 4.4 Hide Text tab in the implemented web interface.", "Normal"),
        ("image", IMG_HIDE_TEXT, "Normal"),
        ("code", "if method == \"LSB\":\n    get_cbs().hide_text_directly(text, cover_path, out_path)\nelif method == \"DCT\":\n    get_dct().hide_text_dct(cover_path, text, out_path)", "Normal"),
        ("text", "Explanation: User text is embedded into the selected cover image using the chosen steganography method.", "Normal"),

        ("text_bold", "Figure 4.5 Hide Barcode tab for embedding pre-generated barcode carriers.", "Normal"),
        ("image", IMG_HIDE_BARCODE, "Normal"),
        ("code", "pattern = get_cbs().load_barcode(barcode_path)\npattern = _scale_barcode_pattern_to_fill_cover(pattern, cover_path, method, jpeg_robust)\nget_barcode_stego().hide_barcode_in_image(cover_path, pattern, out_path)", "Normal"),
        ("text", "Explanation: A pre-generated barcode is loaded, optionally scaled for capacity, then embedded into the cover image.", "Normal"),

        ("text", "4.7 Steganography Embedding, Extraction, and Quality Evaluation", "Title"),
        ("text", "4.7.1 LSB-Based Steganography", "Title"),
        ("text", "LSB provides high embedding capacity with low visible distortion for lossless image paths.", "Normal"),
        ("text", "4.7.2 Data Extraction Process", "Title"),
        ("text", "Extraction reads hidden bits and reconstructs text or barcode payloads.", "Normal"),
        ("text", "4.7.3 DCT-Based Steganography", "Title"),
        ("text", "DCT embedding improves robustness by operating in frequency coefficients.", "Normal"),
        ("text", "4.7.4 Hybrid Embedding Strategy", "Title"),
        ("text", "Hybrid mode combines LSB and DCT to balance quality and robustness.", "Normal"),
        ("text", "4.7.5 Quality Metrics Evaluation", "Title"),
        ("code", "metrics = QualityMetrics.compare_images(cover_img, stego_img)\n# PSNR, SSIM, BER and advanced indicators are returned", "Normal"),
        ("text", "Explanation: Objective metrics are computed to evaluate visual quality and embedding impact.", "Normal"),

        ("text", "4.8 Noise Handling and Denoising", "Title"),
        ("text", "Noise simulation and denoising improve practical reliability analysis.", "Normal"),
        ("code", "if method == \"median\":\n    arr = apply_median_filter(arr, size)\nelse:\n    arr = apply_gaussian_filter(arr, sigma)", "Normal"),
        ("text", "Explanation: Median and Gaussian filters are applied to reduce noise before extraction.", "Normal"),
        ("text", "4.8.1 Effect of DCT Embedding and Noise on Quality Metrics", "Title"),
        ("text", "Metrics are monitored before and after noise/denoising to quantify robustness.", "Normal"),

        ("text", "4.9 Graphical User Interface", "Title"),
        ("text_bold", "Figure 4.6 Extract tab for recovering data from stego images.", "Normal"),
        ("image", IMG_EXTRACT, "Normal"),
        ("code", "if extract_type == \"barcode\":\n    pattern = get_barcode_stego().extract_barcode_from_image(path)\n    text = _decode_barcode_pattern(get_cbs(), pattern, use_hybrid_decryption=True)", "Normal"),
        ("text", "Explanation: The extract flow supports barcode recovery and robust decoding from stego images.", "Normal"),

        ("text", "4.10 End-to-End Execution Flow", "Title"),
        ("text", "4.10.1 Sender Side", "Title"),
        ("text", "Input -> optional encryption -> barcode generation -> embedding -> stego output.", "Normal"),
        ("text", "4.10.2 Receiver Side", "Title"),
        ("text", "Stego input -> extraction -> decode -> optional decryption -> recovered message.", "Normal"),

        ("text", "4.11 Summary", "Title"),
        ("text", "This chapter presented a complete implementation with real successful interface captures, code-level workflow snippets, and concise explanations for each stage.", "Normal"),
    ]

    for kind, value, style in entries:
        if kind == "text":
            insert_text(anchor, value, style=style)
        elif kind == "text_bold":
            insert_text(anchor, value, style=style, bold=True)
        elif kind == "code":
            insert_code(anchor, value)
        elif kind == "image":
            insert_image(anchor, value)

    for p in list(doc.paragraphs[2:ch5_idx]):
        delete_paragraph(p)

    doc.save(str(DOC_PATH))
    print(str(DOC_PATH))
    print(str(BACKUP_PATH))


if __name__ == "__main__":
    main()
