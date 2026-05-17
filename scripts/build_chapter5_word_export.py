"""Build standalone Chapter V (5.1–5.5) as a .docx in stego_project folder."""
from docx import Document

OUT = r"c:\Users\pipis\OneDrive\Desktop\newgrad\stego_project\Chapter_V_System_Testing_5.1-5.5.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def main() -> None:
    doc = Document()
    add_heading(doc, "Chapter V — System Testing and Performance Evaluation", 0)
    add_para(doc, "")

    add_heading(doc, "5.1 Introduction", 2)
    add_para(
        doc,
        "This chapter presents the testing and evaluation of the proposed steganography system. "
        "The objective is to verify that the system meets the functional and non-functional "
        "requirements defined in Chapter 3. The testing process includes automated verification "
        "of the Flask web API (local test client), live deployment checks against the public HTTPS "
        "endpoint (stegax.design), manual and browser-based smoke tests, and performance evaluation "
        "using PSNR, SSIM, and BER where applicable. Field testing on mobile (iPhone) was also used "
        "to validate real-user workflows (save/share stego images, long-running extract operations).",
    )
    add_para(
        doc,
        "Testing covers: (1) barcode encode/decode with optional hybrid (AES+RSA) encryption; "
        "(2) hide text and hide barcode using LSB, DCT, and Hybrid embedding; (3) symmetric "
        "extraction (same method as embedding); (4) quality metrics API; (5) service reliability "
        "under concurrent HTTP traffic when long-running jobs execute (Gunicorn worker model).",
    )

    add_heading(doc, "5.2 Testing Setup", 2)
    add_heading(doc, "5.2.1 Test Environment", 3)
    add_para(
        doc,
        "Development: PC with Windows 10/11, 8–16 GB RAM, SSD; remote SSH to DigitalOcean droplet "
        "(Ubuntu, Python 3.12, venv). Deployment server: cloud VPS at 46.101.244.129, hostname Stegax, "
        "domain stegax.design behind Nginx reverse proxy (TLS) to Gunicorn on 127.0.0.1:5000.",
    )
    add_para(
        doc,
        "Python 3.x; Flask web application; Gunicorn with gthread worker (1 worker, 8 threads) and "
        "600 s timeout; Nginx with extended proxy read/send timeouts and large client body limit for "
        "image uploads; systemd unit stego-web.service for automatic restarts.",
    )
    add_para(
        doc,
        "Pillow, NumPy, SciPy (DCT), scikit-image, cryptography (Flask app requirements as pinned in "
        "requirements.txt). Desktop Tkinter UI optional; primary validation path is the web UI + REST API.",
    )
    add_heading(doc, "5.2.2 Test Data", 3)
    add_para(
        doc,
        "PNG and BMP for lossless barcode carriers and stego I/O; JPEG accepted where the pipeline allows "
        "(DCT path supports JPEG with robust mode). TIFF supported for extract/upload in the web app.",
    )
    add_para(
        doc,
        "Covers from small test grids (e.g. 256×256, 320×320, 384×384) up to larger photos; barcode "
        "pattern scaled to cover dimensions in hide-barcode. Large patterns trigger omission of heavy "
        "JSON previews to protect mobile browsers.",
    )
    add_para(
        doc,
        "Short ASCII strings through longer payloads; hybrid-encrypted barcodes; stress test of repeated "
        "extract-barcode (20 consecutive API calls on the same stego PNG) to detect intermittent failures.",
    )

    add_heading(doc, "5.3 Testing Separate S/W Units", 2)
    add_heading(doc, "5.3.1 Unit Testing", 3)
    add_para(
        doc,
        "Each module is tested independently (automated script scripts/run_app_verification.py — 9/9 checks):",
    )
    add_para(
        doc,
        "AES encryption module — AES stage of the hybrid pipeline exercised via Encode with "
        "“Use hybrid cryptography” and Decode with hybrid decryption enabled in the barcode engine.",
    )
    add_para(
        doc,
        "RSA module — Covered as part of the hybrid encode/decode round-trip (RSA-wrapped AES key material).",
    )
    add_para(
        doc,
        "QR generation module — Custom barcode (QR-like grid) generation and rendering to PNG via /api/encode; "
        "load/decode via /api/decode and pattern-based extract paths.",
    )
    add_para(
        doc,
        "LSB embedding — LSB hide-text and hide-barcode; vectorized bit packing/extraction for performance; "
        "PNG output with controlled compression level for barcode stego.",
    )
    add_para(
        doc,
        "DCT embedding — DCT hide-text and hide-barcode; optional JPEG-robust coefficient quantization; strict "
        "DCTBAR header matching on lossless PNG after field bugfix (prevents false “success” when wrong extract "
        "method is chosen).",
    )
    add_heading(doc, "5.3.2 Integration Testing", 3)
    add_para(
        doc,
        "Form POST /api/encode → PNG barcode download; optional X-Progress-Log header; JSON legacy path "
        "retained with ?format=json for tests.",
    )
    add_para(
        doc,
        "Multipart /api/hide-text and /api/hide-barcode with cover + payload; hide-barcode default response "
        "raw PNG with X-Metrics-Json / X-Progress-Log headers (avoids huge base64-in-JSON failures).",
    )
    add_para(
        doc,
        "Round-trip for each method (LSB, DCT, Hybrid): hide then extract-text or extract-barcode with "
        "matching method; live HTTPS tests to stegax.design confirmed LSB/DCT/Hybrid barcode pipeline.",
    )
    add_heading(doc, "5.3.3 System Testing", 3)
    add_para(
        doc,
        "End-to-end system testing: browser smoke test on stegax.design (Encode → Generate barcode succeeded); "
        "API-level hide-barcode + extract-barcode for all three methods with shared test assets; desktop "
        "verification suite passes in full.",
    )
    add_para(
        doc,
        "Verify correct message recovery: extracted text matches original payload for text and barcode-decode "
        "modes when embedding and extraction methods match; mismatch (e.g. LSB stego extracted as DCT) yields "
        "errors or garbage — documented as user-configuration risk mitigated by UI hints.",
    )
    add_para(
        doc,
        "Verify image quality preservation: PSNR/SSIM/BER from /api/metrics and hide-barcode response headers "
        "when output size/pixel count within server thresholds; skipped for very large outputs to avoid stalls.",
    )

    add_heading(doc, "5.4 Testing Integrated System", 2)
    add_para(
        doc,
        "Full pipeline execution on production stack: client → Nginx (443) → Gunicorn (127.0.0.1:5000) → "
        "Flask app; systemd-managed service restart after code updates; pip install -r requirements.txt in venv.",
    )
    add_para(doc, "Validate:")
    add_para(
        doc,
        "Correct extraction validated for aligned method selection; Decode tab must receive the barcode PNG "
        "from Encode — uploading a stego image to Decode produces “Invalid barcode format” (observed on mobile).",
    )
    add_para(
        doc,
        "Data integrity: lossless PNG roundtrip preserves embedded barcode bits for LSB/DCT/Hybrid when the "
        "image is not recompressed (iPhone: re-saving stego as JPEG can harm DCT unless JPEG-robust hiding was used).",
    )
    add_para(
        doc,
        "System reliability: replaced single synchronous Gunicorn worker with gthread + threads so long "
        "hide/extract jobs no longer block the entire site; extended client, proxy, and JavaScript timeouts "
        "to reduce “failed to fetch” on mobile.",
    )

    add_heading(doc, "5.5 A/B Testing", 2)
    add_para(
        doc,
        "Compare different techniques (LSB vs DCT vs Hybrid) on imperceptibility, robustness, and measured "
        "metrics where available; field issues below are included in the test record.",
    )
    add_para(
        doc,
        "Recorded defects and resolutions during testing: (1) Slow or failing hide-barcode / extract — "
        "mitigated by vectorized NumPy operations in LSB/DCT/Hybrid barcode paths and binary PNG responses. "
        "(2) Browser/proxy timeouts and “failed to fetch” on long jobs — mitigated by Gunicorn --timeout 600, "
        "Nginx proxy timeouts, 10-minute fetch in web JS, gthread worker. (3) Entire site unresponsive during "
        "one long request — mitigated by gthread + multi-threaded worker instead of one blocking sync worker. "
        "(4) iPhone: DCT extract showed garbage while LSB worked — root cause: extraction method must match "
        "embedding method; fuzzy DCT header on lossless images could false-match wrong stego; fixed by exact "
        "header for PNG/BMP/TIFF; UI hints added under Hide Barcode and Extract. (5) Decode tab error on "
        "mobile — users uploaded stego instead of raw barcode image. (6) IDE browser automation could not "
        "set files on hidden <input type=\"file\"> (not visible to Playwright); validation used API instead. "
        "(7) Deployment: no SSH key in cloud IDE — upload must run deploy.ps1 from Windows; on server, "
        "bash-only — attempts to run .ps1 or cd C:\\... failed until Linux commands (systemctl restart "
        "stego-web) were used. (8) On Linux servers use curl, not curl.exe.",
    )
    add_para(
        doc,
        "Automated regression: scripts/run_app_verification.py reports 9/9 passed (encode/decode, "
        "hide/extract text LSB+DCT+Hybrid, hide/extract barcode LSB+DCT+Hybrid, 20× LSB extract stress, "
        "metrics API).",
    )

    tbl = doc.add_table(rows=4, cols=5)
    tbl.style = "Table Grid"
    hdr = ["Method", "PSNR", "SSIM", "BER", "Robustness"]
    for c, h in enumerate(hdr):
        tbl.rows[0].cells[c].text = h
    data = [
        (
            "LSB",
            "Typically highest among the three on lossless PNG (minimal pixel change).",
            "High structural similarity to cover when capacity allows.",
            "Very low BER on lossless roundtrip; rises if image is recompressed or edited.",
            "Weakest to lossy JPEG/social recompression; best for pristine PNG workflow.",
        ),
        (
            "DCT",
            "Moderate; coefficient quantization trades imperceptibility for robustness options.",
            "Good when blocks align; JPEG-robust mode targets recompression survival.",
            "Low on lossless roundtrip; normal DCT hiding is sensitive to JPEG if not using JPEG-robust.",
            "Better than pure LSB against JPEG when JPEG-robust embedding is enabled.",
        ),
        (
            "Hybrid",
            "Intermediate—splits payload between spatial LSB and frequency DCT channels.",
            "Generally good; reflects combined spatial/frequency perturbation.",
            "Low on lossless roundtrip with matched extraction; method must match embed.",
            "Combines aspects of LSB and DCT; user must keep extract method aligned with embed.",
        ),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            tbl.rows[r].cells[c].text = val

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
