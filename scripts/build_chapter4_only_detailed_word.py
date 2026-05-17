from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


OUT = r"C:\Users\pipis\Downloads\CSE_GradII_Chapter4_Only_Detailed_Final.docx"

IMG_HOME = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-home.png"
IMG_ENCODE_OK = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-encode-success.png"
IMG_DECODE = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-decode-tab.png"
IMG_HIDE_TEXT = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-hide-text-tab.png"
IMG_HIDE_BARCODE = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-hide-barcode-tab.png"
IMG_EXTRACT = r"C:\Users\pipis\AppData\Local\Temp\cursor\screenshots\stegax-extract-tab.png"


def qn(tag):
    prefix, local = tag.split(":")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return "{%s}%s" % (ns[prefix], local)


def add_code(doc, text):
    p = doc.add_paragraph("")
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F3F3")
    p._element.get_or_add_pPr().append(shd)


def add_fig(doc, caption, path):
    doc.add_paragraph(caption).runs[0].bold = True
    p = doc.add_paragraph("")
    p.add_run().add_picture(path, width=Inches(5.9))


def main():
    d = Document()
    d.add_paragraph("CHAPTER IV", style="Title")
    d.add_paragraph("Preliminary Implementation", style="Title")

    d.add_paragraph("4.1 Development Environment", style="Heading 1")
    d.add_paragraph(
        "The system was implemented using Python 3 in a modular architecture that separates core steganography logic, barcode processing, cryptographic components, and the web interface layer. "
        "The development workflow relied on Flask for API orchestration, NumPy for matrix manipulation, and Pillow for image loading and conversion. "
        "This combination enabled rapid iteration while preserving deterministic behavior across encode, hide, extract, and decode operations."
    )
    d.add_paragraph(
        "For runtime deployment, the web service was configured behind Gunicorn and managed by systemd to ensure continuous operation. "
        "This production setup was selected to support stable 24/7 access and to reduce manual intervention during demonstrations. "
        "The resulting deployment model is reproducible and can be migrated to other cloud instances with minimal configuration changes."
    )
    add_code(
        d,
        'app = Flask(__name__, template_folder="templates", static_folder="static")\n'
        'app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024  # 200 MB\n'
        "gunicorn -w 2 -b 127.0.0.1:5000 'web.app:app' --chdir ."
    )

    d.add_paragraph("4.2 System Architecture Implementation", style="Heading 1")
    d.add_paragraph(
        "The implemented architecture follows a layered and modular design. The web layer handles user interactions and API validation, "
        "the core layer implements embedding and extraction algorithms (LSB, DCT, Hybrid), and support layers provide barcode transformation, "
        "cryptographic protection, and objective quality metrics. This structure simplifies debugging and allows each subsystem to be tested independently."
    )
    d.add_paragraph(
        "During operation, the pipeline proceeds through deterministic stages: input normalization, optional encryption, barcode/payload formatting, embedding, "
        "stego image generation, extraction, and decoding. This sequencing makes the system predictable in runtime behavior and suitable for formal evaluation."
    )
    add_code(
        d,
        "from core import BarcodeSteganography, DCTSteganography, HybridSteganography\n"
        "from barcode import CustomBarcodeSteganography\n"
        "from metrics import QualityMetrics"
    )

    d.add_paragraph("4.3 Encryption Modules", style="Heading 1")
    d.add_paragraph("4.3.1 SecureEncryption Module", style="Heading 2")
    d.add_paragraph(
        "The SecureEncryption path protects payload confidentiality before any embedding stage. "
        "Instead of inserting raw plaintext directly, the module transforms content into a protected form that is less interpretable even if recovered by unauthorized extraction. "
        "This design adds a security layer that complements, rather than replaces, steganographic concealment."
    )
    add_code(
        d,
        "if use_hybrid_encryption and self.hybrid_crypto:\n"
        "    encrypted_result = self.hybrid_crypto.encrypt_text_hybrid(text)\n"
        "    data_to_encode = '[HYBRID_ENCRYPTED]' + encrypted_result"
    )
    d.add_paragraph("4.3.2 HybridCrypto Module", style="Heading 2")
    d.add_paragraph(
        "The HybridCrypto module combines symmetric and asymmetric cryptographic principles. "
        "Symmetric processing is used for practical payload encryption throughput, while RSA key material is used for secure key handling. "
        "This hybrid model was selected to balance computational efficiency and secure key exchange requirements."
    )
    add_code(
        d,
        "private_key, public_key = HybridCrypto.load_rsa_keys()\n"
        "self.hybrid_crypto = HybridCrypto(rsa_public_key=public_key, rsa_private_key=private_key)"
    )

    d.add_paragraph("4.4 Custom Barcode Carrier Generation", style="Heading 1")
    d.add_paragraph(
        "The system converts prepared payloads into a structured barcode-like carrier before embedding into cover images. "
        "This representation creates a stable intermediate format that can be hidden, extracted, and decoded with explicit header and payload boundaries. "
        "Carrier generation is deterministic and designed to support both plaintext and encrypted payload streams."
    )
    add_fig(d, "Figure 4.1 Main application interface of the steganography system.", IMG_HOME)
    add_code(
        d,
        "@app.route('/api/encode', methods=['POST'])\n"
        "def api_encode():\n"
        "    pattern = get_cbs().encode(text, use_hybrid_encryption=use_hybrid)\n"
        "    return send_file(_barcode_pattern_to_bytes(pattern), mimetype='image/png')"
    )
    d.add_paragraph("Brief: Figure 4.1 shows the full web interface and the main operational tabs used throughout the implementation workflow.")
    add_fig(d, "Figure 4.2 Encode tab showing successful barcode generation and download.", IMG_ENCODE_OK)
    d.add_paragraph(
        "The successful encode state confirms that payload parsing, pattern generation, and PNG rendering completed correctly. "
        "This stage is critical because all downstream hide/extract operations depend on the structural correctness of the generated carrier."
    )
    d.add_paragraph("Brief: Figure 4.2 confirms a successful encode transaction with progress logs and completion status.")

    d.add_paragraph("4.5 Barcode Decoding", style="Heading 1")
    d.add_paragraph(
        "The decode stage reverses carrier generation by loading a barcode image, recovering the encoded bit sequence, validating headers, "
        "and reconstructing payload text. The implementation includes defensive checks to detect invalid formats and prevent silent corruption."
    )
    add_fig(d, "Figure 4.3 Decode tab for barcode recovery workflow.", IMG_DECODE)
    add_code(
        d,
        "pattern = cbs.load_barcode(path)\n"
        "text = cbs.decode(pattern, use_hybrid_decryption=True)\n"
        "return jsonify({'text': text, 'is_image': False})"
    )
    d.add_paragraph("Brief: Figure 4.3 illustrates the decode interface used to recover hidden payload from barcode images.")

    d.add_paragraph("4.6 Hiding Existing Barcode (Embedding Pre-Generated Barcode)", style="Heading 1")
    d.add_paragraph("4.6.1 Workflow", style="Heading 2")
    d.add_paragraph(
        "In this mode, the user supplies a previously generated barcode and a cover image. The application reads the barcode pattern, "
        "optionally scales it according to cover capacity, and performs embedding using the selected algorithm. "
        "This flow is useful in evaluation scenarios where carrier generation and embedding are tested as separate steps."
    )
    add_fig(d, "Figure 4.4 Hide Text tab in the implemented web interface.", IMG_HIDE_TEXT)
    add_code(
        d,
        "if method == 'LSB':\n"
        "    get_cbs().hide_text_directly(text, cover_path, out_path)\n"
        "elif method == 'DCT':\n"
        "    get_dct().hide_text_dct(cover_path, text, out_path)\n"
        "else:\n"
        "    get_hybrid().hide_text_hybrid(cover_path, text, out_path)"
    )
    d.add_paragraph("Brief: Figure 4.4 presents the hide-text operation where plaintext is embedded into a selected cover image.")
    add_fig(d, "Figure 4.5 Hide Barcode tab for embedding pre-generated barcode carriers.", IMG_HIDE_BARCODE)
    add_code(
        d,
        "pattern = get_cbs().load_barcode(barcode_path)\n"
        "pattern = _scale_barcode_pattern_to_fill_cover(pattern, cover_path, method, jpeg_robust)\n"
        "get_barcode_stego().hide_barcode_in_image(cover_path, pattern, out_path)"
    )
    d.add_paragraph("Brief: Figure 4.5 shows the barcode-embedding mode for inserting a pre-generated carrier into a cover image.")

    d.add_paragraph("4.7 Steganography Embedding, Extraction, and Quality Evaluation", style="Heading 1")
    d.add_paragraph("4.7.1 LSB-Based Steganography", style="Heading 2")
    d.add_paragraph(
        "LSB embedding provides high payload capacity and excellent visual imperceptibility for lossless image paths. "
        "It is used as the primary baseline due to low complexity and efficient execution time."
    )
    d.add_paragraph("4.7.2 Data Extraction Process", style="Heading 2")
    d.add_paragraph(
        "Extraction routines recover hidden data from stego images and route payloads to the appropriate decode path. "
        "Recent reliability improvements include robust JSON parsing in the client and additional fallback logic in barcode decode paths."
    )
    d.add_paragraph("4.7.3 DCT-Based Steganography", style="Heading 2")
    d.add_paragraph(
        "DCT embedding modifies selected frequency coefficients rather than direct pixel LSBs. "
        "This strategy improves robustness against some transformations while preserving acceptable visual quality."
    )
    d.add_paragraph("4.7.4 Hybrid Embedding Strategy", style="Heading 2")
    d.add_paragraph(
        "Hybrid embedding combines spatial and frequency-domain channels to balance robustness and capacity. "
        "The implementation uses controlled bit distribution between LSB and DCT paths and includes extraction logic that reconstructs the combined payload reliably."
    )
    d.add_paragraph("4.7.5 Quality Metrics Evaluation", style="Heading 2")
    d.add_paragraph(
        "Objective metrics are calculated to quantify distortion and extraction quality. "
        "Reported values include PSNR, SSIM, and BER, together with auxiliary analytical indicators used in the UI comparison panel."
    )
    add_code(
        d,
        "metrics = QualityMetrics.compare_images(cover_img, stego_img)\n"
        "# PSNR, SSIM, BER and advanced indicators"
    )

    d.add_paragraph("4.8 Graphical User Interface", style="Heading 1")
    d.add_paragraph(
        "The GUI is organized as operational tabs: Encode, Decode, Hide Text, Hide Barcode, and Extract. "
        "Each tab includes progress feedback, operation logs, and clear status messages to improve usability during demonstrations and evaluation sessions. "
        "The interface also includes inline previews for extracted barcode/image outputs."
    )
    add_fig(d, "Figure 4.6 Extract tab for recovering data from stego images.", IMG_EXTRACT)
    add_code(
        d,
        "if extract_type == 'barcode':\n"
        "    pattern = get_barcode_stego().extract_barcode_from_image(path)\n"
        "    text = _decode_barcode_pattern(get_cbs(), pattern, use_hybrid_decryption=True)"
    )
    d.add_paragraph("Brief: Figure 4.6 demonstrates the extraction interface used to recover and decode hidden data from stego images.")

    d.add_paragraph("4.9 End-to-End Execution Flow", style="Heading 1")
    d.add_paragraph("4.9.1 Sender Side", style="Heading 2")
    d.add_paragraph(
        "Sender workflow: input preparation -> optional hybrid encryption -> barcode generation -> embedding into selected cover image -> stego output export. "
        "This pipeline ensures that both concealment and confidentiality are applied before transmission."
    )
    d.add_paragraph("4.9.2 Receiver Side", style="Heading 2")
    d.add_paragraph(
        "Receiver workflow: stego upload -> extraction -> barcode/text decoding -> optional decryption -> final content reconstruction. "
        "Error handling and fallback logic were enhanced to minimize intermittent extraction failures under practical usage."
    )

    d.add_paragraph("4.10 Summary", style="Heading 1")
    d.add_paragraph(
        "This chapter presented a full implementation of the proposed secure steganography system, including architecture, module-level behavior, "
        "carrier generation, embedding/extraction methods, and GUI-level interaction flow. The implementation is no longer conceptual: "
        "it is deployed and validated through repeated runtime tests across all major operational paths."
    )
    d.add_paragraph(
        "Recent stabilization work resolved practical runtime issues observed during real usage, including intermittent extraction failures, "
        "hybrid-path data type errors, and API serialization issues in metrics output. Verification testing confirmed reliable behavior across "
        "Encode/Decode, Hide/Extract Text (LSB, DCT, Hybrid), Hide/Extract Barcode (LSB, DCT, Hybrid), and repeated extraction stress scenarios."
    )
    d.add_paragraph(
        "Accordingly, the implemented system demonstrates functional completeness, operational stability, and user-facing clarity suitable for academic demonstration. "
        "The next chapter can therefore focus on formal testing analysis and performance interpretation rather than implementation uncertainty."
    )

    d.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

